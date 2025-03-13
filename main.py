import PySimpleGUI as sg
from docx import Document
import sys
import os
import platform
import subprocess

def resource_path(relative_path):
    """
    PyInstaller実行時には、_MEIPASSディレクトリからファイルを参照するためのヘルパー関数。
    """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def insert_image_to_template(image_path, template_path, output_path):
    doc = Document(template_path)
    table = doc.tables[0]
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.add_run().add_picture(image_path)
    doc.save(output_path)

def open_folder(folder_path):
    if platform.system() == "Windows":
        os.startfile(folder_path)
    elif platform.system() == "Darwin":
        subprocess.Popen(["open", folder_path])
    else:
        subprocess.Popen(["xdg-open", folder_path])

def main():
    layout = [
        [sg.Text("誰の名刺ですか？"),sg.I("", key="-name-",size=(30, 1))],
        [sg.Text("挿入するJPG画像を選択してください"),sg.FileBrowse("画像を選択", file_types=(("JPEG Files", "*.jpg"),), target="-FILE-")],
        [sg.Input(key="-FILE-", visible=False)],
        [sg.Button("生成"), sg.Button("フォルダを表示"), sg.Button("終了")]
    ]
    
    window = sg.Window("10card_create", layout, font = ("HGP教科書体", 15))
    
    while True:
        event, values = window.read()
        if event in (None, "終了"):
            break
        
        if event == "生成":
            image_path = values["-FILE-"]
            name = values["-name-"]
            if not image_path:
                sg.popup("画像ファイルを選択してください")
                continue
            elif not name:
                sg.popup("対象者の名前を入力してください")
                continue
            
            try:
                template_path = resource_path("card_template.docx")
                os.makedirs("card_docx", exist_ok=True)
                output_path = f"card_docx/{name}-名刺.docx"
                insert_image_to_template(image_path, template_path, output_path)
                sg.popup("名刺用紙が生成されました", f"出力ファイル：{output_path}")
                window["-name-"].update("")
                window["-FILE-"].update("")
            except Exception as e:
                sg.popup("エラーが発生しました", str(e))
        elif event == "フォルダを表示":
            folder_path = "card_docx"
            os.makedirs("card_docx", exist_ok=True)
            open_folder(folder_path)
    
    window.close()

if __name__ == "__main__":
    main()
